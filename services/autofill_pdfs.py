from datetime import datetime
import os
from docx import Document
import pytz

def nome_pdf(nome_do_andidato):
    nome_do_andidato = (str(nome_do_andidato).upper()).replace(" ", "_")
    today = datetime.now(pytz.timezone('Brazil/East'))
    data_nova = str(today.strftime("_%Y_%m_%d"))
    contract = "Formulário_de_contrato".upper()
    return (contract + '_' + nome_do_andidato + '_' + data_nova + ".docx")

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

def form_filler_for_contract(dados, out_name):
    path_pasta_upload = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        'website',
        'uploads',
        out_name
    ).replace('\\', '/')

    parent_directory = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    in_path = os.path.join(parent_directory, "doc_base", "base_contract.docx")
    text_string = ""
    sign_string = ""
    child_string = ""
    children_info_string = ""
    n = 1

    for person in dados:
        text_string += f"{person['fullName']}, aqui denominado CONTRATANTE, brasileiro, portador da Carteira de Identidade (RG) n {person['rg']}, titular do CPF n. {person['cpf']}, residente e domiciliado à {person['address']}\n"
        sign_string += f"\n_____________________________\n{person['fullName']}\n\n"
        for child in person['children']:
            child_string += f"{n}. Filho do sr. {person['fullName']}: {child['childName']}\n\n"
            n += 1
        for children in person['children']:
            children_info_string += f"{n}. {children['childName']}, CPF {children['child_cpf']}. , data de nascimento {children['child_dob']}\n\n"
    brazil_east_timezone = pytz.timezone('America/Sao_Paulo')
    today_date = datetime.now(brazil_east_timezone).date()

    portuguese_month_names = [
        "janeiro", "fevereiro", "março", "abril", "maio", "junho",
        "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"
    ]
    formatted_date = f"{today_date.day} de {portuguese_month_names[today_date.month - 1]} de {today_date.year}"

    variables = {
        "Person_details": text_string,
        "Persons_signs": sign_string,
        "Child_details": child_string,
        "children_info": children_info_string,
        "today_date": str(formatted_date)
    }

    template_document = Document(in_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(path_pasta_upload)

    return True
